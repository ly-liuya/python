from docx import Document
from docx.shared import Cm,Pt   # 用于设置大小（图片）
from docx.document import Document as doc

# 创建代表word文档的Document对象
document = Document()
# 添加标题
document.add_heading("快快乐乐学习python",level=1)

# 添加段落
p = document.add_paragraph("python是一门非常流行的语言")
run = p.add_run("非常棒")
run.bold = True   # 字体加粗
run.font.size = Pt(18)  # 设置字体大小
run.underline = False  # 是否设置下划线
p.add_run("!")

# 保存文档
document.save("./data/demo.docx")

