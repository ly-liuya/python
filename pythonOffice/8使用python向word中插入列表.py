from docx import Document

# 列表分为有序和无序列表

# 创建word文档对象
document = Document()

# 添加无序列表
document.add_paragraph("第一步：hello",style="List Bullet")

document.add_paragraph("第二步：word",style="List Bullet")

# 添加有序列表
document.add_paragraph("刘备",style="List Number")
document.add_paragraph("关羽",style="List Number")
document.add_paragraph("张飞",style="List Number")

document.save("./data/列表.docx")


